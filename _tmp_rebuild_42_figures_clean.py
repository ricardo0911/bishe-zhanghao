from __future__ import annotations

import os
import shutil
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


def U(s: str) -> str:
    return s.encode("ascii").decode("unicode_escape")


FIG = U(r"\u56fe")  # 图
SEE = U(r"\u5982\u56fe")  # 如图
SHOWN = U(r"\u6240\u793a\u3002")  # 所示。
FONT_SONG = U(r"\u5b8b\u4f53")
FONT_HEI = U(r"\u9ed1\u4f53")


def has_image_para(p: Paragraph) -> bool:
    xml = p._p.xml
    return ("w:drawing" in xml) or ("v:imagedata" in xml)


def para_index(doc: Document, para: Paragraph) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p._p is para._p:
            return i
    return -1


def find_heading(doc: Document, prefix: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if (p.text or "").strip().startswith(prefix):
            return i
    return -1


def remove_para(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)


def clear_runs(p: Paragraph) -> None:
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_para_text(
    p: Paragraph,
    text: str,
    font_name: str,
    size_pt: float,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    clear_runs(p)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if align is not None:
        p.alignment = align


def insert_before(base: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    base._p.addprevious(new_p)
    return Paragraph(new_p, base._parent)


def insert_after(base: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    base._p.addnext(new_p)
    return Paragraph(new_p, base._parent)


def clone_image_para(src_para: Paragraph, after_para: Paragraph) -> Paragraph:
    new_elem = deepcopy(src_para._p)
    after_para._p.addnext(new_elem)
    return Paragraph(new_elem, after_para._parent)


def is_target_ref_or_caption(t: str) -> bool:
    s = (t or "").strip()
    if not s:
        return False
    if s.startswith("图4-"):
        return True
    if ("图4-" in s) and ("如图4-" in s) and ("所示" in s):
        return True
    return False


def is_caption_text(s: str) -> bool:
    t = (s or "").strip()
    if not t.startswith(FIG):
        return False
    tail = t[1:].strip().replace(U(r"\uff0d"), "-").replace(" ", "")
    if "-" not in tail:
        return False
    left, right = tail.split("-", 1)
    if not left.isdigit():
        return False
    right_digits = ""
    for ch in right:
        if ch.isdigit():
            right_digits += ch
        else:
            break
    return bool(right_digits)


def normalize_images(doc: Document) -> None:
    for p in doc.paragraphs:
        if has_image_para(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_w = Cm(15)
    for shp in doc.inline_shapes:
        try:
            if shp.width and shp.width > max_w:
                shp.width = max_w
        except Exception:
            pass


def enforce_caption_style(doc: Document) -> None:
    for p in doc.paragraphs:
        if is_caption_text(p.text or ""):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not p.runs:
                p.add_run(p.text or "")
            for run in p.runs:
                run.font.name = FONT_HEI
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEI)
                run.font.size = Pt(10.5)
                run.font.bold = False


def main() -> None:
    target = os.environ.get("DOC_TARGET", "").strip()
    if not target:
        raise SystemExit("DOC_TARGET is required")
    path = Path(target)
    if not path.exists():
        raise SystemExit(f"not found: {path}")

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = path.with_name(path.stem + f".bak_rebuild_{stamp}.docx")
    shutil.copy2(path, backup)

    doc = Document(str(path))

    sec_ids = ["4.2.1.1", "4.2.1.2", "4.2.1.3", "4.2.2.1", "4.2.2.2", "4.2.2.3"]
    sec_starts = {sid: find_heading(doc, sid) for sid in sec_ids}
    if any(v < 0 for v in sec_starts.values()):
        missing = [k for k, v in sec_starts.items() if v < 0]
        raise SystemExit(f"missing headings: {missing}")

    block_start = sec_starts["4.2.1.1"]
    block_end = find_heading(doc, "4.2.3")
    if block_end < 0:
        block_end = len(doc.paragraphs)

    # 1) clear all previous 4-9.. references/captions in 4.2.1.1~4.2.2.3
    for i in range(block_end - 1, block_start - 1, -1):
        t = (doc.paragraphs[i].text or "").strip()
        if is_target_ref_or_caption(t):
            remove_para(doc.paragraphs[i])

    # recompute starts after deletion
    sec_starts = {sid: find_heading(doc, sid) for sid in sec_ids}

    title_map = {
        "4.2.1.1": [
            U(r"\u652f\u4ed8\u529f\u80fd\u622a\u56fe"),
            U(r"\u6295\u4fdd\u8ba2\u5355\u786e\u8ba4\u622a\u56fe"),
            U(r"\u652f\u4ed8\u7ed3\u679c\u9875\u9762\u622a\u56fe"),
        ],
        "4.2.1.2": [
            U(r"\u7406\u8d54\u7533\u8bf7\u529f\u80fd\u622a\u56fe"),
            U(r"\u7406\u8d54\u8fdb\u5ea6\u67e5\u8be2\u529f\u80fd\u622a\u56fe"),
            U(r"\u7406\u8d54\u8be6\u60c5\u67e5\u770b\u622a\u56fe"),
        ],
        "4.2.1.3": [
            U(r"\u6551\u63f4\u7533\u8bf7\u529f\u80fd\u622a\u56fe"),
            U(r"\u6551\u63f4\u8bb0\u5f55\u67e5\u8be2\u622a\u56fe"),
            U(r"\u8f85\u52a9\u670d\u52a1\u529f\u80fd\u622a\u56fe"),
        ],
        "4.2.2.1": [
            U(r"\u8ba2\u5355\u5ba1\u6838\u5904\u7406\u622a\u56fe"),
            U(r"\u8ba2\u5355\u72b6\u6001\u6d41\u8f6c\u622a\u56fe"),
            U(r"\u8ba2\u5355\u652f\u4ed8\u8054\u52a8\u622a\u56fe"),
        ],
        "4.2.2.2": [
            U(r"\u7406\u8d54\u521d\u5ba1\u64cd\u4f5c\u622a\u56fe"),
            U(r"\u7406\u8d54\u8fdb\u5ea6\u8ddf\u8e2a\u622a\u56fe"),
            U(r"\u7406\u8d54\u5ba1\u6838\u7ed3\u679c\u622a\u56fe"),
        ],
        "4.2.2.3": [
            U(r"\u8f66\u8f86\u4fe1\u606f\u6838\u9a8c\u622a\u56fe"),
            U(r"\u6551\u63f4\u534f\u540c\u8c03\u5ea6\u622a\u56fe"),
            U(r"\u534f\u540c\u5904\u7406\u7ed3\u679c\u622a\u56fe"),
        ],
    }

    fig_no = 9
    summary: list[str] = []

    for i, sid in enumerate(sec_ids):
        st = find_heading(doc, sid)
        next_sid = sec_ids[i + 1] if i + 1 < len(sec_ids) else None
        if next_sid:
            ed = find_heading(doc, next_sid)
            if ed < 0:
                ed = len(doc.paragraphs)
        else:
            ed = find_heading(doc, "4.2.3")
            if ed < 0:
                ed = len(doc.paragraphs)

        section_paras = doc.paragraphs[st:ed]
        imgs = [p for p in section_paras if has_image_para(p)]
        if not imgs:
            # if none, borrow nearest image around section
            src = None
            around = []
            for k, p in enumerate(doc.paragraphs):
                if has_image_para(p):
                    around.append((abs(k - st), p))
            if around:
                around.sort(key=lambda x: x[0])
                src = around[0][1]
            if src is None:
                continue
            anchor = doc.paragraphs[st]
            img = clone_image_para(src, anchor)
            img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            imgs = [img]

        while len(imgs) < 3:
            new_img = clone_image_para(imgs[-1], imgs[-1])
            new_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            imgs.append(new_img)

        if len(imgs) > 3:
            for extra in imgs[3:]:
                remove_para(extra)
            imgs = imgs[:3]

        titles = title_map[sid]
        for j, img in enumerate(imgs):
            fig = f"4-{fig_no}"
            title = titles[j]
            mention = f"{title}{SEE}{fig}{SHOWN}"
            if sid == "4.2.1.1" and j == 0:
                mention = U(r"\u652f\u4ed8\u529f\u80fd\u622a\u56fe\u5982\u56fe4-9\u6240\u793a\u3002")
            caption = f"{FIG}{fig} {title}"

            prev = insert_before(img)
            set_para_text(prev, mention, FONT_SONG, 12)

            cap = insert_after(img)
            set_para_text(cap, caption, FONT_HEI, 10.5, WD_ALIGN_PARAGRAPH.CENTER)

            summary.append(f"{sid} -> {fig} {title}")
            fig_no += 1

    normalize_images(doc)
    enforce_caption_style(doc)
    doc.save(str(path))

    print(f"UPDATED={path}")
    print(f"BACKUP={backup}")
    print(f"TOTAL={len(summary)}")
    if summary:
        print(f"RANGE=4-9..4-{8 + len(summary)}")
    for line in summary:
        print(line)


if __name__ == "__main__":
    main()

