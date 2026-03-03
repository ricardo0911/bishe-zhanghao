from __future__ import annotations

import os
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


def U(text: str) -> str:
    return text.encode("ascii").decode("unicode_escape")


FIG_CHAR = U(r"\u56fe")
SEE_CHAR = U(r"\u5982\u56fe")
SHOWN_CHAR = U(r"\u6240\u793a\u3002")
FONT_SONG = U(r"\u5b8b\u4f53")
FONT_HEI = U(r"\u9ed1\u4f53")


def has_image_para(p: Paragraph) -> bool:
    xml = p._p.xml
    return ("w:drawing" in xml) or ("v:imagedata" in xml)


def para_index(doc: Document, para: Paragraph) -> int:
    for idx, p in enumerate(doc.paragraphs):
        if p._p is para._p:
            return idx
    return -1


def clear_runs(p: Paragraph) -> None:
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_para_text(
    p: Paragraph,
    text: str,
    font_name: str,
    size_pt: float,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    clear_runs(p)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if align is not None:
        p.alignment = align


def insert_before(base: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    base._p.addprevious(new_p)
    return Paragraph(new_p, base._parent)


def insert_after(base: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    base._p.addnext(new_p)
    return Paragraph(new_p, base._parent)


def remove_para(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)


def is_caption_text(text: str) -> bool:
    t = (text or "").strip()
    if not t.startswith(FIG_CHAR):
        return False
    s = t[1:].strip().replace(U(r"\uff0d"), "-").replace(" ", "")
    if "-" not in s:
        return False
    left, right = s.split("-", 1)
    if not left.isdigit():
        return False
    right_digits = ""
    for ch in right:
        if ch.isdigit():
            right_digits += ch
        else:
            break
    return bool(right_digits)


def clone_image_para(src_para: Paragraph, after_para: Paragraph) -> Paragraph:
    new_elem = deepcopy(src_para._p)
    after_para._p.addnext(new_elem)
    return Paragraph(new_elem, after_para._parent)


def find_heading_idx(doc: Document, starts_with: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t.startswith(starts_with):
            return i
    return -1


def get_section_range(doc: Document, sid: str, next_sid: str | None) -> tuple[int, int]:
    st = find_heading_idx(doc, sid)
    if st < 0:
        return -1, -1
    if next_sid:
        ed = find_heading_idx(doc, next_sid)
        if ed < 0:
            ed = len(doc.paragraphs)
    else:
        ed = len(doc.paragraphs)
        for i in range(st + 1, len(doc.paragraphs)):
            t = (doc.paragraphs[i].text or "").strip()
            if t.startswith("4.2.3") or t.startswith("5 "):
                ed = i
                break
    return st, ed


def nearest_image_para(doc: Document, target_idx: int) -> Paragraph | None:
    imgs: list[tuple[int, Paragraph]] = []
    for i, p in enumerate(doc.paragraphs):
        if has_image_para(p):
            imgs.append((i, p))
    if not imgs:
        return None
    imgs.sort(key=lambda x: abs(x[0] - target_idx))
    return imgs[0][1]


def enforce_all_images_centered_and_resized(doc: Document) -> None:
    for p in doc.paragraphs:
        if has_image_para(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_width = Cm(15)
    for shape in doc.inline_shapes:
        try:
            if shape.width and shape.width > max_width:
                shape.width = max_width
        except Exception:
            pass


def enforce_caption_style(p: Paragraph) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not p.runs:
        p.add_run(p.text or "")
    for run in p.runs:
        run.font.name = FONT_HEI
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEI)
        run.font.size = Pt(10.5)
        run.font.bold = False


def main() -> None:
    target = os.environ.get("DOC_TARGET", "").strip()
    if not target:
        raise SystemExit("DOC_TARGET env is required")

    path = Path(target)
    if not path.exists():
        raise SystemExit(f"doc not found: {path}")

    doc = Document(str(path))
    section_ids = [
        "4.2.1.1",
        "4.2.1.2",
        "4.2.1.3",
        "4.2.2.1",
        "4.2.2.2",
        "4.2.2.3",
    ]

    title_map: dict[str, list[str]] = {
        "4.2.1.1": [
            U(r"\u652f\u4ed8\u529f\u80fd\u9875\u9762\u622a\u56fe"),
            U(r"\u6295\u4fdd\u8ba2\u5355\u786e\u8ba4\u622a\u56fe"),
            U(r"\u652f\u4ed8\u7ed3\u679c\u9875\u9762\u622a\u56fe"),
        ],
        "4.2.1.2": [
            U(r"\u7406\u8d54\u7533\u8bf7\u63d0\u4ea4\u622a\u56fe"),
            U(r"\u7406\u8d54\u8fdb\u5ea6\u5217\u8868\u622a\u56fe"),
            U(r"\u7406\u8d54\u8be6\u60c5\u67e5\u770b\u622a\u56fe"),
        ],
        "4.2.1.3": [
            U(r"\u6551\u63f4\u7533\u8bf7\u63d0\u4ea4\u622a\u56fe"),
            U(r"\u6551\u63f4\u8bb0\u5f55\u67e5\u8be2\u622a\u56fe"),
            U(r"\u8f85\u52a9\u670d\u52a1\u529f\u80fd\u622a\u56fe"),
        ],
        "4.2.2.1": [
            U(r"\u8ba2\u5355\u5ba1\u6838\u5904\u7406\u622a\u56fe"),
            U(r"\u8ba2\u5355\u72b6\u6001\u6d41\u8f6c\u622a\u56fe"),
            U(r"\u8ba2\u5355\u652f\u4ed8\u8054\u52a8\u622a\u56fe"),
        ],
        "4.2.2.2": [
            U(r"\u7406\u8d54\u521d\u5ba1\u64cd\u4f5c\u622a\u56fe"),
            U(r"\u7406\u8d54\u8fdb\u5ea6\u8ddf\u8e2a\u622a\u56fe"),
            U(r"\u7406\u8d54\u5ba1\u6838\u7ed3\u679c\u622a\u56fe"),
        ],
        "4.2.2.3": [
            U(r"\u8f66\u8f86\u4fe1\u606f\u6838\u9a8c\u622a\u56fe"),
            U(r"\u6551\u63f4\u534f\u540c\u8c03\u5ea6\u622a\u56fe"),
            U(r"\u534f\u540c\u5904\u7406\u7ed3\u679c\u622a\u56fe"),
        ],
    }

    fig_no = 9
    summary: list[str] = []

    for i, sid in enumerate(section_ids):
        next_sid = section_ids[i + 1] if i + 1 < len(section_ids) else None
        st, ed = get_section_range(doc, sid, next_sid)
        if st < 0:
            continue

        section_paras = doc.paragraphs[st:ed]
        imgs = [p for p in section_paras if has_image_para(p)]

        # fallback if section has no image
        if not imgs:
            src = nearest_image_para(doc, st)
            if src is None:
                continue
            anchor = doc.paragraphs[st]
            new_img = clone_image_para(src, anchor)
            new_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            imgs = [new_img]

        while len(imgs) < 3:
            new_img = clone_image_para(imgs[-1], imgs[-1])
            new_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            imgs.append(new_img)

        if len(imgs) > 3:
            for extra in imgs[3:]:
                remove_para(extra)
            imgs = imgs[:3]

        titles = title_map[sid]
        for j, img_para in enumerate(imgs):
            fig = f"4-{fig_no}"
            title = titles[j]
            mention = f"{title}{SEE_CHAR}{fig}{SHOWN_CHAR}"
            if sid == "4.2.1.1" and j == 0:
                mention = U(r"\u652f\u4ed8\u529f\u80fd\u622a\u56fe\u5982\u56fe4-9\u6240\u793a\u3002")
            caption = f"{FIG_CHAR}{fig} {title}"

            idx = para_index(doc, img_para)
            prev = doc.paragraphs[idx - 1] if idx > 0 else None
            if prev is None:
                prev = insert_before(img_para)
                set_para_text(prev, mention, FONT_SONG, 12)
            else:
                prev_text = (prev.text or "").strip()
                prev_is_heading = prev_text.startswith("4.") or prev_text.startswith("5 ")
                if has_image_para(prev) or prev_is_heading or prev_text.startswith(FIG_CHAR) or (SEE_CHAR in prev_text):
                    new_prev = insert_before(img_para)
                    set_para_text(new_prev, mention, FONT_SONG, 12)
                elif (
                    (not prev_text)
                    or ("?" in prev_text)
                ):
                    set_para_text(prev, mention, FONT_SONG, 12)
                else:
                    new_prev = insert_before(img_para)
                    set_para_text(new_prev, mention, FONT_SONG, 12)

            idx = para_index(doc, img_para)
            nxt = doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else None
            if nxt is None:
                nxt = insert_after(img_para)
                set_para_text(nxt, caption, FONT_HEI, 10.5, WD_ALIGN_PARAGRAPH.CENTER)
            else:
                next_text = (nxt.text or "").strip()
                next_is_heading = next_text.startswith("4.") or next_text.startswith("5 ")
                if has_image_para(nxt) or next_is_heading:
                    new_nxt = insert_after(img_para)
                    set_para_text(new_nxt, caption, FONT_HEI, 10.5, WD_ALIGN_PARAGRAPH.CENTER)
                elif is_caption_text(next_text) or (not next_text) or ("?" in next_text):
                    set_para_text(nxt, caption, FONT_HEI, 10.5, WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    new_nxt = insert_after(img_para)
                    set_para_text(new_nxt, caption, FONT_HEI, 10.5, WD_ALIGN_PARAGRAPH.CENTER)

            summary.append(f"{sid} -> {fig} {title}")
            fig_no += 1

    enforce_all_images_centered_and_resized(doc)
    for p in doc.paragraphs:
        if is_caption_text((p.text or "").strip()):
            enforce_caption_style(p)

    doc.save(str(path))

    print(f"UPDATED={path}")
    print(f"TOTAL_ASSIGNED={len(summary)}")
    if summary:
        print(f"FIG_RANGE=4-9..4-{8 + len(summary)}")
    for line in summary:
        print(line)


if __name__ == "__main__":
    main()
